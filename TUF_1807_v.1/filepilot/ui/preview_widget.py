"""
ui/preview_widget.py
Area di anteprima grande (sinistra). Mostra immagini, video (con
riproduzione vera, non solo un frame), la prima pagina dei PDF, un
estratto testuale dei documenti Word e un segnaposto per i modelli 3D
(STL/OBJ) e i formati GoPro/360/DJI non riproducibili direttamente.
Precarica il file immagine successivo per prestazioni elevate.
Supporta lo zoom con la rotellina del mouse sulle immagini e mostra
un contatore "x / y" in basso a destra.
"""
from __future__ import annotations

from pathlib import Path

from PySide6.QtCore import Qt, QSize, QUrl, QThread, Signal
from PySide6.QtGui import QPixmap, QImage, QDesktopServices, QPainter, QPainterPath
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QTextEdit, QStackedWidget,
    QSizePolicy, QScrollArea, QPushButton, QSlider,
)

try:
    from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
    from PySide6.QtMultimediaWidgets import QVideoWidget
    _HAS_MULTIMEDIA = True
except ImportError:
    _HAS_MULTIMEDIA = False

from filepilot.models import FileItem, FileCategory

# RICHIESTA: "riesci ad implementare con le icone giuste degli zip/
# musica ecc?" — l'icona mostrata in alto a destra nel visualizzatore,
# per i tipi NON multimediali, usava emoji generiche. Ora usa le
# icone disegnate apposta per TUF (filepilot/assets/icons/<ext>.png,
# una per estensione, non solo per macroarea), estendendo la
# copertura a TUTTE le categorie senza un'anteprima visiva vera
# (prima mancavano Audio, Ebook, Sottotitoli, Testo/dati, Font,
# Applicazioni). Foto/Video/GoPro-360-Drone non la mostrano, dato che
# gia' si vede (o si prova a mostrare) il contenuto vero.
_NO_BADGE_CATEGORIES = {FileCategory.IMAGE, FileCategory.VIDEO, FileCategory.ACTION}

_BADGE_ASSET_DIR = Path(__file__).resolve().parent.parent / "assets" / "icons"
_badge_pixmap_cache: dict[str, QPixmap | None] = {}


def _load_badge_pixmap(ext: str) -> QPixmap | None:
    """Carica (con cache in memoria, MAI su disco: vedi nota 'no log'
    sulle altre cache in questo stesso file) l'icona per una
    estensione. Ogni icona e' un file statico incluso con l'app
    (filepilot/assets/icons/), non generato ne' scaricato a runtime."""
    key = ext.lstrip(".").lower()
    if key in _badge_pixmap_cache:
        return _badge_pixmap_cache[key]
    path = _BADGE_ASSET_DIR / f"{key}.png"
    pix = QPixmap(str(path)) if path.exists() else None
    if pix is not None and pix.isNull():
        pix = None
    _badge_pixmap_cache[key] = pix
    return pix


def _rounded_pixmap(pix: QPixmap, size: int, radius: int) -> QPixmap:
    scaled = pix.scaled(size, size, Qt.IgnoreAspectRatio, Qt.SmoothTransformation)
    result = QPixmap(size, size)
    result.fill(Qt.transparent)
    painter = QPainter(result)
    painter.setRenderHint(QPainter.Antialiasing)
    clip = QPainterPath()
    clip.addRoundedRect(0, 0, size, size, radius, radius)
    painter.setClipPath(clip)
    painter.drawPixmap(0, 0, scaled)
    painter.end()
    return result

try:
    import cv2
except ImportError:
    cv2 = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx as python_docx
except ImportError:
    python_docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


RENDER_MAX = QSize(2200, 2200)  # risoluzione "base" per permettere zoom in senza sgranare troppo


LAST_RENDER_ERROR: str | None = None


def render_pixmap(item: FileItem, max_size: QSize = RENDER_MAX) -> QPixmap | None:
    """Genera una QPixmap di anteprima ad alta risoluzione (usata poi
    per l'adattamento a schermo e per lo zoom). Per i video genera un
    frame di riferimento, usato come miniatura (es. per l'animazione
    di spostamento), non per la riproduzione vera e propria."""
    global LAST_RENDER_ERROR
    cat = item.category
    path = item.path
    try:
        if cat == FileCategory.IMAGE:
            pix = QPixmap(path)
            if pix.isNull():
                return None
            if pix.width() > max_size.width() or pix.height() > max_size.height():
                pix = pix.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            return pix

        if cat == FileCategory.VIDEO and cv2 is not None:
            cap = cv2.VideoCapture(path)
            ok, frame = cap.read()
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
            if ok and frame_count > 10:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_count // 3)
                ok2, mid_frame = cap.read()
                if ok2:
                    frame = mid_frame
            cap.release()
            if not ok:
                return None
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            h, w, ch = frame.shape
            qimg = QImage(frame.data, w, h, ch * w, QImage.Format_RGB888)
            pix = QPixmap.fromImage(qimg.copy())
            return pix.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)

        if cat == FileCategory.PDF:
            if fitz is None:
                LAST_RENDER_ERROR = "PyMuPDF (fitz) non installato"
                return None
            doc = fitz.open(path)
            page = doc.load_page(0)
            zoom = 2.5
            mat = fitz.Matrix(zoom, zoom)
            pm = page.get_pixmap(matrix=mat)
            fmt = QImage.Format_RGBA8888 if pm.alpha else QImage.Format_RGB888
            qimg = QImage(pm.samples, pm.width, pm.height, pm.stride, fmt)
            pix = QPixmap.fromImage(qimg.copy())
            doc.close()
            return pix.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
    except Exception as e:
        import traceback
        traceback.print_exc()  # visibile nel terminale, utile per capire cosa non va
        LAST_RENDER_ERROR = f"{type(e).__name__}: {e}"
        return None
    return None


def _decode_to_qimage(item: FileItem, max_size: QSize) -> QImage | None:
    """Come render_pixmap, ma per foto e PDF ritorna un QImage invece
    di una QPixmap: QImage e' sicuro da creare e manipolare FUORI dal
    thread dell'interfaccia (QPixmap invece no, va sempre creata/usata
    nel thread principale). Usato dal worker in background qui sotto,
    cosi' un file lento da decodificare (una foto enorme, un PDF con
    contenuti complessi) non blocca mai l'interfaccia: prima capitava
    ordinando la lista per dimensione, perche' il file piu' grande
    diventava quello corrente e veniva decodificato in modo sincrono
    sul thread principale."""
    global LAST_RENDER_ERROR
    cat = item.category
    path = item.path
    try:
        if cat == FileCategory.IMAGE:
            img = QImage(path)
            if img.isNull():
                return None
            if img.width() > max_size.width() or img.height() > max_size.height():
                img = img.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            return img

        if cat == FileCategory.VIDEO and cv2 is not None:
            cap = cv2.VideoCapture(path)
            ok, frame = cap.read()
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
            if ok and frame_count > 10:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_count // 3)
                ok2, mid_frame = cap.read()
                if ok2:
                    frame = mid_frame
            cap.release()
            if not ok:
                return None
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            h, w, ch = frame.shape
            img = QImage(frame.data, w, h, ch * w, QImage.Format_RGB888).copy()
            return img.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)

        if cat == FileCategory.PDF:
            if fitz is None:
                LAST_RENDER_ERROR = "PyMuPDF (fitz) non installato"
                return None
            doc = fitz.open(path)
            page = doc.load_page(0)
            zoom = 2.5
            mat = fitz.Matrix(zoom, zoom)
            pm = page.get_pixmap(matrix=mat)
            fmt = QImage.Format_RGBA8888 if pm.alpha else QImage.Format_RGB888
            img = QImage(pm.samples, pm.width, pm.height, pm.stride, fmt).copy()
            doc.close()
            return img.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
    except Exception as e:
        import traceback
        traceback.print_exc()
        LAST_RENDER_ERROR = f"{type(e).__name__}: {e}"
        return None
    return None


class _PreviewRenderWorker(QThread):
    """Decodifica foto/PDF in un thread separato (vedi _decode_to_qimage
    sopra sul perche' un QImage e non una QPixmap). Usato sia per
    mostrare il file corrente senza bloccare l'interfaccia, sia per
    precaricare in anticipo quello successivo."""

    done = Signal(str, object)  # (path del file, QImage oppure None)

    def __init__(self, item: FileItem, max_size: QSize, parent=None):
        super().__init__(parent)
        self._item = item
        self._max_size = max_size

    def run(self) -> None:
        img = _decode_to_qimage(self._item, self._max_size)
        self.done.emit(self._item.path, img)


def get_video_duration_seconds(path: str) -> float | None:
    """Calcola la durata di un video in secondi (frame_count / fps).
    Usata nella finestra di controllo duplicati per mostrare la durata
    accanto a nome/dimensione/date. Ritorna None se non calcolabile."""
    if cv2 is None:
        return None
    try:
        cap = cv2.VideoCapture(path)
        fps = cap.get(cv2.CAP_PROP_FPS) or 0
        frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0
        cap.release()
        if fps > 0 and frame_count > 0:
            return frame_count / fps
    except Exception:
        pass
    return None


def get_video_thumb_and_duration(
    path: str, max_size: QSize = RENDER_MAX
) -> tuple[QPixmap | None, float | None]:
    """Apre il video UNA SOLA VOLTA per ottenere sia la miniatura (un
    fotogramma centrale, con la stessa logica di render_pixmap) sia la
    durata (frame_count / fps). Usata nella finestra di controllo
    duplicati: prima si apriva il video due volte (una per la
    miniatura, una per la durata) — con tanti video nello stesso
    controllo duplicati il doppio lavoro sincrono poteva rallentare
    parecchio l'apertura della finestra (rischio di "non risponde")."""
    if cv2 is None:
        return None, None
    try:
        cap = cv2.VideoCapture(path)
        fps = cap.get(cv2.CAP_PROP_FPS) or 0
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        duration = frame_count / fps if fps > 0 and frame_count > 0 else None

        ok, frame = cap.read()
        if ok and frame_count > 10:
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_count // 3)
            ok2, mid_frame = cap.read()
            if ok2:
                frame = mid_frame
        cap.release()
        if not ok:
            return None, duration
        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        h, w, ch = frame.shape
        qimg = QImage(frame.data, w, h, ch * w, QImage.Format_RGB888)
        pix = QPixmap.fromImage(qimg.copy())
        return pix.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation), duration
    except Exception:
        return None, None

def decode_thumb_and_duration(item: FileItem, max_size: QSize) -> tuple[QImage | None, float | None]:
    """Come get_video_thumb_and_duration, ma ritorna un QImage invece
    di una QPixmap (quindi e' sicuro chiamarla FUORI dal thread
    dell'interfaccia, vedi _decode_to_qimage) e funziona per QUALSIASI
    categoria supportata, non solo i video. Usata dal caricamento
    asincrono delle miniature nella finestra duplicati (vedi
    ui/duplicate_review_dialog.py): con lotti fino a 100 file,
    decodificarle in modo sincrono sul thread principale poteva
    bloccare l'interfaccia per parecchi secondi, specialmente
    cambiando l'ordinamento (che riparte sempre da miniature non
    ancora in cache)."""
    if item.category == FileCategory.VIDEO:
        if cv2 is None:
            return None, None
        try:
            cap = cv2.VideoCapture(item.path)
            fps = cap.get(cv2.CAP_PROP_FPS) or 0
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
            duration = frame_count / fps if fps > 0 and frame_count > 0 else None

            ok, frame = cap.read()
            if ok and frame_count > 10:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_count // 3)
                ok2, mid_frame = cap.read()
                if ok2:
                    frame = mid_frame
            cap.release()
            if not ok:
                return None, duration
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            h, w, ch = frame.shape
            img = QImage(frame.data, w, h, ch * w, QImage.Format_RGB888).copy()
            return img.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation), duration
        except Exception:
            return None, None
    return _decode_to_qimage(item, max_size), None


def extract_docx_text(path: str, max_chars: int = 4000) -> str:
    if python_docx is None:
        return "Anteprima Word non disponibile (python-docx mancante)."
    try:
        doc = python_docx.Document(path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paras)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[...]"
        return text or "(documento vuoto)"
    except Exception as e:
        return f"Impossibile leggere il documento: {e}"


def _esc(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;"))


def extract_docx_html(path: str, max_chars: int = 12000) -> str:
    """RICHIESTA ("visualizzatore di Word ed Excel"): prima l'anteprima
    Word era solo testo grezzo (extract_docx_text sopra, ancora usata
    come fallback), senza titoli in evidenza ne' tabelle. Qui si
    costruisce un HTML semplice (titoli in grassetto/piu' grandi in
    base allo stile del paragrafo, tabelle rese come vere tabelle HTML)
    da mostrare nello stesso QTextEdit gia' usato per il testo, che
    supporta un sottoinsieme di HTML.

    NB: python-docx espone paragrafi e tabelle come DUE collezioni
    separate (doc.paragraphs, doc.tables) e non l'ordine ESATTO in cui
    si alternano nel documento originale (richiederebbe di scorrere
    l'XML grezzo del corpo del documento) — qui si mostrano prima tutti
    i paragrafi e poi tutte le tabelle. Per un'anteprima va bene,
    l'ordine preciso conta relativamente poco."""
    if python_docx is None:
        return "<p>Anteprima Word non disponibile (python-docx mancante).</p>"
    try:
        doc = python_docx.Document(path)
        parts: list[str] = []
        total_chars = 0

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name or "").lower() if p.style else ""
            if "heading 1" in style_name or "title" in style_name:
                parts.append(f"<h2>{_esc(text)}</h2>")
            elif "heading" in style_name:
                parts.append(f"<h3>{_esc(text)}</h3>")
            else:
                parts.append(f"<p>{_esc(text)}</p>")
            total_chars += len(text)
            if total_chars > max_chars:
                parts.append("<p>[...]</p>")
                break

        for table in doc.tables:
            rows_html = []
            for row in table.rows:
                cells_html = "".join(f"<td>{_esc(c.text)}</td>" for c in row.cells)
                rows_html.append(f"<tr>{cells_html}</tr>")
            if rows_html:
                parts.append(
                    "<table border='1' cellspacing='0' cellpadding='4' "
                    "style='border-color:#444; margin-top:10px;'>"
                    + "".join(rows_html) + "</table>"
                )

        if not parts:
            return "<p>(documento vuoto)</p>"
        return "".join(parts)
    except Exception as e:
        return f"<p>Impossibile leggere il documento: {_esc(str(e))}</p>"


def extract_xlsx_html(path: str, max_rows: int = 200, max_cols: int = 30) -> str:
    """RICHIESTA ("visualizzatore di Word ed Excel"): prima Excel non
    mostrava NESSUN contenuto, solo un'icona segnaposto con un
    pulsante "Apri con...". Qui si legge il primo foglio con openpyxl
    e lo si mostra come una vera tabella HTML, cosi' come per Word nel
    QTextEdit gia' esistente. Se il file ha piu' fogli, viene mostrato
    solo il primo (con una nota che dice quanti altri ce ne sono) —
    aggiungere un selettore di foglio e' possibile in futuro ma non
    necessario per una prima anteprima utile."""
    if openpyxl is None:
        return "<p>Anteprima Excel non disponibile (libreria openpyxl mancante).</p>"
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheet = wb.worksheets[0]
        parts = [f"<h3>{_esc(sheet.title)}</h3>"]
        if len(wb.sheetnames) > 1:
            others = ", ".join(wb.sheetnames[1:])
            parts.append(f"<p style='color:#888;'>Altri fogli non mostrati: {_esc(others)}</p>")

        # BUG RISOLTO (in fase di test): iterare sempre fino a max_cols
        # (30) anche per fogli che usano solo 2-3 colonne riempiva la
        # tabella di celle vuote inutili, rendendo l'anteprima molto
        # piu' larga e vuota del necessario. Ora ci si ferma alle
        # colonne/righe EFFETTIVAMENTE usate dal foglio (sheet.max_row/
        # max_column), con max_rows/max_cols solo come tetto massimo
        # per fogli davvero enormi.
        used_rows = min(sheet.max_row or 0, max_rows)
        used_cols = min(sheet.max_column or 0, max_cols)
        rows_html = []
        row_count = 0
        for row in sheet.iter_rows(max_row=used_rows, max_col=used_cols):
            cells_html = []
            for cell in row:
                value = cell.value
                text = "" if value is None else str(value)
                cells_html.append(f"<td>{_esc(text)}</td>")
            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
            row_count += 1
        wb.close()

        if not rows_html:
            return "".join(parts) + "<p>(foglio vuoto)</p>"

        parts.append(
            "<table border='1' cellspacing='0' cellpadding='4' "
            "style='border-color:#444;'>" + "".join(rows_html) + "</table>"
        )
        if sheet.max_row and sheet.max_row > max_rows:
            parts.append(f"<p style='color:#888;'>[...e altre {sheet.max_row - max_rows} righe]</p>")
        return "".join(parts)
    except Exception as e:
        return f"<p>Impossibile leggere il foglio: {_esc(str(e))}</p>"


def extract_code_text(path: str, max_chars: int = 6000) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(max_chars + 1)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[...]"
        return text or "(file vuoto)"
    except OSError as e:
        return f"Impossibile leggere il file: {e}"


def extract_zip_listing(path: str, max_entries: int = 300) -> str:
    import zipfile
    try:
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
        lines = [f"Archivio ZIP — {len(names)} elementi\n"]
        for name in names[:max_entries]:
            lines.append(name)
        if len(names) > max_entries:
            lines.append(f"\n[...e altri {len(names) - max_entries} elementi]")
        return "\n".join(lines)
    except Exception as e:
        return f"Impossibile leggere l'archivio: {e}"


class ClickableSlider(QSlider):
    """QSlider normale, ma un clic diretto sulla barra porta subito il
    cursore in quel punto (comportamento di default di Qt: un clic
    sposta solo di un piccolo passo, serve trascinare)."""

    def _value_from_x(self, x: float) -> int:
        if self.maximum() <= self.minimum():
            return self.minimum()
        ratio = min(max(x / max(self.width(), 1), 0.0), 1.0)
        return int(self.minimum() + ratio * (self.maximum() - self.minimum()))

    def mousePressEvent(self, event) -> None:
        if event.button() == Qt.LeftButton:
            value = self._value_from_x(event.position().x())
            self.setValue(value)
            self.sliderMoved.emit(value)
            event.accept()
        else:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event) -> None:
        if event.buttons() & Qt.LeftButton:
            value = self._value_from_x(event.position().x())
            self.setValue(value)
            self.sliderMoved.emit(value)
            event.accept()
        else:
            super().mouseMoveEvent(event)


def _format_ms(ms: int) -> str:
    total_sec = max(0, ms) // 1000
    m, s = divmod(total_sec, 60)
    return f"{m:02d}:{s:02d}"


class _PosterWithPlayButton(QWidget):
    """Miniatura del video + tasto play grande al centro, entrambi
    widget "normali" di Qt (non il video vero). Tenerli su una pagina
    separata dal QVideoWidget evita un problema noto di Qt: il
    QVideoWidget usa una superficie di rendering nativa che spesso
    disegna SOPRA qualunque widget sovrapposto, anche se in teoria e'
    piu' in basso nello z-order. Con due pagine distinte non serve
    sovrapporre nulla al video vero."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.poster_label = QLabel(self)
        self.poster_label.setAlignment(Qt.AlignCenter)
        self.poster_label.setScaledContents(True)
        self.poster_label.setStyleSheet("background-color: #000;")

        self.play_button = QPushButton("▶", self)
        self.play_button.setFixedSize(84, 84)
        self.play_button.setStyleSheet(
            "QPushButton { background-color: rgba(0,0,0,140); color: #fff;"
            " border: 3px solid #fff; border-radius: 42px; font-size: 30px; }"
            "QPushButton:hover { background-color: rgba(0,0,0,190); }"
        )

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.poster_label.setGeometry(0, 0, self.width(), self.height())
        bw, bh = self.play_button.width(), self.play_button.height()
        self.play_button.move((self.width() - bw) // 2, (self.height() - bh) // 2)


class PreviewWidget(QWidget):
    """Pannello grande di anteprima, al centro dell'area sinistra."""

    ZOOM_MIN = 0.2
    ZOOM_MAX = 6.0
    ZOOM_STEP = 1.15

    def __init__(self, parent=None):
        super().__init__(parent)
        # NO LOG: cache SOLO in memoria di processo (mai su disco),
        # sparisce alla chiusura del programma — vedi il commento in
        # cima a config.py sull'unico elenco di cose che TUF salva
        # davvero.
        self._cache: dict[str, QPixmap] = {}
        self._current_item: FileItem | None = None
        self._base_pixmap: QPixmap | None = None
        self._zoom = 1.0  # 1.0 = adatta alla finestra
        self._render_worker: _PreviewRenderWorker | None = None
        self._preload_workers: list[_PreviewRenderWorker] = []
        self._pending_render_item: FileItem | None = None
        self._pending_render_kind: str | None = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        self.stack = QStackedWidget(self)

        # --- pagina immagini/pdf, con zoom ---
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setAlignment(Qt.AlignCenter)
        self.scroll.setStyleSheet("background-color: #1a1a1a; border: none;")

        self.image_label = QLabel("Nessun file selezionato")
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.image_label.setStyleSheet("background-color: #1a1a1a; color: #888;")
        self.scroll.setWidget(self.image_label)
        self.scroll.wheelEvent = self._on_scroll_wheel  # cosi' la rotellina zooma sempre, anche da scrollata

        # --- pagina video, con riproduzione vera ---
        self.video_page = QWidget()
        video_layout = QVBoxLayout(self.video_page)
        video_layout.setContentsMargins(0, 0, 0, 0)
        video_layout.setSpacing(0)

        self.player = None
        self.audio_output = None
        if _HAS_MULTIMEDIA:
            self.video_widget = QVideoWidget()
            self.video_widget.setStyleSheet("background-color: #000;")

            # pagina "poster": miniatura + tasto play grande, mostrata
            # finche' non si preme play. Al primo play si passa alla
            # pagina col video vero (vedi nota nella classe qui sopra
            # sul perche' non sono sovrapposti).
            self.video_poster_widget = _PosterWithPlayButton()
            self.video_poster_label = self.video_poster_widget.poster_label
            self.big_play_btn = self.video_poster_widget.play_button
            self.big_play_btn.clicked.connect(self._on_big_play_clicked)

            self.video_inner_stack = QStackedWidget()
            self.video_inner_stack.addWidget(self.video_poster_widget)  # index 0: poster+play
            self.video_inner_stack.addWidget(self.video_widget)         # index 1: video vero

            self.player = QMediaPlayer(self)
            self.audio_output = QAudioOutput(self)
            self.player.setAudioOutput(self.audio_output)
            self.player.setVideoOutput(self.video_widget)
            self.player.positionChanged.connect(self._on_video_position_changed)
            self.player.durationChanged.connect(self._on_video_duration_changed)
            self.player.playbackStateChanged.connect(self._on_video_state_changed)
            video_layout.addWidget(self.video_inner_stack, stretch=1)

            controls = QHBoxLayout()
            controls.setContentsMargins(8, 6, 8, 6)
            self.play_btn = QPushButton("▶")
            self.play_btn.setFixedWidth(40)
            self.play_btn.clicked.connect(self._toggle_play_pause)
            controls.addWidget(self.play_btn)

            self.video_time_label = QLabel("00:00 / 00:00")
            self.video_time_label.setStyleSheet("color: #ccc; font-size: 11px;")
            controls.addWidget(self.video_time_label)

            self.video_slider = ClickableSlider(Qt.Horizontal)
            self.video_slider.setRange(0, 0)
            self.video_slider.sliderMoved.connect(self.player.setPosition)
            controls.addWidget(self.video_slider, stretch=1)

            self.speed_btn = QPushButton("1x")
            self.speed_btn.setFixedWidth(44)
            self.speed_btn.setToolTip("Velocità di riproduzione (clic per cambiare)")
            self.speed_btn.clicked.connect(self._cycle_playback_speed)
            controls.addWidget(self.speed_btn)

            video_layout.addLayout(controls)
        else:
            fallback = QLabel(
                "Riproduzione video non disponibile:\ninstalla il modulo QtMultimedia di PySide6."
            )
            fallback.setAlignment(Qt.AlignCenter)
            fallback.setStyleSheet("color: #999;")
            video_layout.addWidget(fallback)

        # --- pagina testo Word ---
        self.text_view = QTextEdit()
        self.text_view.setReadOnly(True)
        self.text_view.setStyleSheet(
            "background-color: #1e1e1e; color: #ddd; font-size: 13px; padding: 16px;"
        )

        # --- pagina segnaposto (modelli 3D, formati GoPro/360/DJI) ---
        self.model_page = QWidget()
        mp_layout = QVBoxLayout(self.model_page)
        mp_layout.setAlignment(Qt.AlignCenter)
        self.model_label = QLabel("🧊\nModello 3D")
        self.model_label.setAlignment(Qt.AlignCenter)
        self.model_label.setStyleSheet("color: #ccc; font-size: 20px;")
        self.model_open_btn = QPushButton("Apri con il visualizzatore predefinito")
        self.model_open_btn.setFixedWidth(280)
        mp_layout.addWidget(self.model_label)
        mp_layout.addWidget(self.model_open_btn, alignment=Qt.AlignCenter)

        self.stack.addWidget(self.scroll)       # index 0: immagini/pdf
        self.stack.addWidget(self.text_view)    # index 1: word
        self.stack.addWidget(self.model_page)   # index 2: segnaposto
        self.stack.addWidget(self.video_page)   # index 3: video

        layout.addWidget(self.stack)

        # --- contatore "x / y" in basso a destra, sovrapposto al riquadro ---
        self.counter_label = QLabel("", self)
        self.counter_label.setStyleSheet(
            "background-color: rgba(0,0,0,150); color: #fff; font-size: 12px;"
            " padding: 4px 8px; border-radius: 4px;"
        )
        self.counter_label.adjustSize()
        self.counter_label.raise_()

        # --- icona tipo file, in alto a destra, per i formati non
        # multimediali (documenti/archivi/codice/3D) ---
        self.type_badge_label = QLabel("", self)
        self.type_badge_label.setAlignment(Qt.AlignCenter)
        self.type_badge_label.setStyleSheet(
            "background-color: rgba(0,0,0,160); color: #fff; border-radius: 8px;"
        )
        self.type_badge_label.hide()

        # --- pulsante "ricentra" in basso a sinistra, visibile solo con zoom attivo ---
        self.recenter_btn = QPushButton("⤢", self)
        self.recenter_btn.setToolTip("Ripristina la vista (Ctrl+0)")
        self.recenter_btn.setFixedSize(34, 34)
        self.recenter_btn.setStyleSheet(
            "QPushButton { background-color: rgba(0,0,0,160); color: #fff;"
            " border-radius: 17px; font-size: 16px; border: none; }"
            "QPushButton:hover { background-color: rgba(0,0,0,210); }"
        )
        self.recenter_btn.clicked.connect(self.reset_zoom)
        self.recenter_btn.hide()

    def sizeHint(self) -> QSize:
        return QSize(1000, 700)

    # ------------------------------------------------------------ API
    def set_counter(self, current: int, total: int) -> None:
        self.counter_label.setText(f"{current} / {total}")
        self.counter_label.adjustSize()
        self._position_counter()
        self.counter_label.raise_()

    @staticmethod
    def format_file_info(item: FileItem) -> str:
        from datetime import datetime
        size = item.size
        for unit in ("B", "KB", "MB", "GB"):
            if size < 1024:
                size_str = f"{size:.1f} {unit}"
                break
            size /= 1024
        else:
            size_str = f"{size:.1f} TB"
        date_str = datetime.fromtimestamp(item.mtime).strftime("%d/%m/%Y %H:%M")
        ext_str = item.extension.upper().lstrip(".")
        return f"{item.name}  •  {size_str}  •  {date_str}  •  .{ext_str}"

    def stop_video(self) -> None:
        """Ferma la riproduzione e rilascia il file video. Va chiamato
        SEMPRE prima di spostare o eliminare un video, altrimenti
        Windows puo' bloccare l'operazione perche' il file e' aperto."""
        if self.player is not None:
            self.player.stop()
            self.player.setSource(QUrl())

    def pause_for_dialog(self) -> None:
        """BUG RISOLTO ("il video resta sopra la finestra di conferma
        eliminazione e nasconde i pulsanti, si esce solo con Esc"): va
        chiamato prima di mostrare QUALUNQUE overlay/dialogo sopra
        l'anteprima. Mette in pausa e torna alla pagina "poster+play"
        (widget Qt normali), perche' il QVideoWidget vero disegna
        SEMPRE sopra gli altri widget anche se e' piu' in basso nello
        z-order (vedi nota in _PosterWithPlayButton) — quindi qualsiasi
        overlay mostrato mentre la pagina video vera e' visibile
        risulterebbe coperto e i suoi pulsanti non cliccabili."""
        if self.player is not None and self.is_showing_video():
            self.player.pause()
            self.video_inner_stack.setCurrentIndex(0)

    def show_item(self, item: FileItem | None) -> None:
        # cambiando file, ferma sempre un eventuale video in riproduzione
        self.stop_video()

        self._current_item = item
        self._zoom = 1.0
        self._update_type_badge(item)
        self.recenter_btn.hide()

        if item is None:
            self.stack.setCurrentIndex(0)
            self._base_pixmap = None
            self.image_label.setText("Nessun file selezionato")
            self.image_label.setPixmap(QPixmap())
            return

        if item.category == FileCategory.VIDEO:
            if self.player is not None:
                self.stack.setCurrentIndex(3)
                self.video_inner_stack.setCurrentIndex(0)  # riparti sempre dalla pagina poster+play
                self.video_poster_label.clear()
                self._start_async_video_poster(item)
                self.player.setSource(QUrl.fromLocalFile(item.path))
                self.player.setPlaybackRate(1.0)
                self.speed_btn.setText("1x")
                self.play_btn.setText("▶")
            else:
                # QtMultimedia non disponibile: torna al vecchio comportamento
                # (frame singolo), anche questo ora caricato in background
                self.stack.setCurrentIndex(0)
                self._base_pixmap = None
                self.image_label.setText("Caricamento anteprima…")
                self.image_label.setPixmap(QPixmap())
                self._start_async_render(item)
            return

        if item.category == FileCategory.WORD:
            # RICHIESTA ("visualizzatore di Word"): prima solo testo
            # grezzo (setPlainText); ora HTML con titoli in evidenza e
            # tabelle vere (vedi extract_docx_html). setPlainText
            # restava come fallback ma extract_docx_html gestisce gia'
            # da sola i casi di errore/libreria mancante, restituendo
            # un messaggio HTML leggibile invece di sollevare.
            self.stack.setCurrentIndex(1)
            self.text_view.setHtml(extract_docx_html(item.path))
            return

        if item.category == FileCategory.EXCEL:
            # RICHIESTA ("visualizzatore di Excel"): prima Excel non
            # mostrava alcun contenuto (rientrava nel gruppo
            # segnaposto qui sotto insieme a modelli 3D/GoPro). Ora ha
            # una vera anteprima tabellare del primo foglio (vedi
            # extract_xlsx_html), mostrata nello stesso QTextEdit usato
            # per Word/codice/archivi.
            self.stack.setCurrentIndex(1)
            self.text_view.setHtml(extract_xlsx_html(item.path))
            return

        if item.category == FileCategory.CODE:
            self.stack.setCurrentIndex(1)
            self.text_view.setPlainText(extract_code_text(item.path))
            return

        if item.category == FileCategory.ARCHIVE:
            self.stack.setCurrentIndex(1)
            self.text_view.setPlainText(extract_zip_listing(item.path))
            return

        if item.category in (FileCategory.MODEL, FileCategory.ACTION,
                              FileCategory.POWERPOINT):
            self.stack.setCurrentIndex(2)
            icons = {
                FileCategory.MODEL: "🧊",
                FileCategory.ACTION: "🎥",
                FileCategory.POWERPOINT: "📽️",
            }
            labels = {
                FileCategory.MODEL: "Modello 3D",
                FileCategory.ACTION: "Formato GoPro / 360 / Drone",
                FileCategory.POWERPOINT: "Presentazione PowerPoint",
            }
            icon = icons[item.category]
            label = labels[item.category]
            self.model_label.setText(f"{icon}\n{label}\n{Path(item.path).name}\n{item.size / 1024:.0f} KB")
            try:
                self.model_open_btn.clicked.disconnect()
            except TypeError:
                pass
            self.model_open_btn.clicked.connect(
                lambda: QDesktopServices.openUrl(QUrl.fromLocalFile(item.path))
            )
            return

        self.stack.setCurrentIndex(0)
        cached = self._cache.pop(item.path, None)
        if cached is not None:
            self._base_pixmap = cached
            self._render_zoom()
        else:
            # BUG RISOLTO ("non risponde" ordinando per dimensione):
            # decodificare qui, in modo sincrono, il file corrente
            # bloccava l'interfaccia finche' non finiva — se il file
            # piu' grande della cartella (es. un PDF pesante, capita
            # spesso ordinando per dimensione) era lento da decodificare,
            # TUF sembrava bloccato per lungo tempo. Ora la decodifica
            # avviene in un thread separato: subito si vede "Caricamento
            # anteprima...", e quando e' pronta compare senza mai aver
            # bloccato il resto dell'interfaccia nel frattempo.
            self._base_pixmap = None
            self.image_label.setText("Caricamento anteprima…")
            self.image_label.setPixmap(QPixmap())
            self._start_async_render(item)

    def _start_async_render(self, item: FileItem) -> None:
        self._request_async_render(item, "photo")

    def _start_async_video_poster(self, item: FileItem) -> None:
        self._request_async_render(item, "video_poster")

    def _request_async_render(self, item: FileItem, kind: str) -> None:
        """Punto unico da cui parte SEMPRE la decodifica in background,
        sia per foto/PDF sia per il poster dei video.
        BUG RISOLTO: prima ogni chiamata avviava subito un thread nuovo;
        navigando molto velocemente tra tanti file grandi (es. tenendo
        premuto 'Avanti', o scorrendo una lista ordinata per dimensione
        piena di file pesanti) si potevano accumulare parecchi thread
        di decodifica insieme, e la CPU tornava a essere sovraccarica
        esattamente come nel problema che dovevano risolvere. Ora c'e'
        sempre AL MASSIMO un thread di decodifica attivo: se arriva una
        nuova richiesta mentre uno e' gia' in corso, si mette in coda
        (sostituendo un'eventuale richiesta precedente ancora in attesa,
        cosi' non si accumulano film intermedi mai piu' visti) e parte
        automaticamente appena quello attuale finisce."""
        if self._render_worker is not None and self._render_worker.isRunning():
            self._pending_render_item = item
            self._pending_render_kind = kind
            return
        self._launch_render_worker(item, kind)

    def _launch_render_worker(self, item: FileItem, kind: str) -> None:
        worker = _PreviewRenderWorker(item, RENDER_MAX, self)
        worker.done.connect(lambda path, qimage, k=kind: self._on_async_render_finished(path, qimage, k))
        self._render_worker = worker
        worker.start()

    def _on_async_render_finished(self, path: str, qimage, kind: str) -> None:
        if kind == "video_poster":
            self._apply_video_poster_result(path, qimage)
        else:
            self._apply_photo_render_result(path, qimage)

        # se nel frattempo e' arrivata una richiesta piu' recente
        # (l'utente e' gia' passato a un altro file), parte subito
        if self._pending_render_item is not None:
            next_item = self._pending_render_item
            next_kind = self._pending_render_kind
            self._pending_render_item = None
            self._pending_render_kind = None
            self._launch_render_worker(next_item, next_kind)

    def _apply_video_poster_result(self, path: str, qimage) -> None:
        if self._current_item is None or self._current_item.path != path:
            return  # file gia' cambiato nel frattempo: risultato scartato
        if self._current_item.category != FileCategory.VIDEO:
            return
        if qimage is not None:
            self.video_poster_label.setPixmap(QPixmap.fromImage(qimage))

    def _apply_photo_render_result(self, path: str, qimage) -> None:
        # se nel frattempo l'utente e' passato a un altro file, questo
        # risultato e' ormai vecchio: si scarta
        if self._current_item is None or self._current_item.path != path:
            return
        item = self._current_item
        pix = QPixmap.fromImage(qimage) if qimage is not None else None
        self._base_pixmap = pix
        if pix is None:
            if item.category == FileCategory.PDF:
                detail = LAST_RENDER_ERROR or "motivo sconosciuto"
                self.image_label.setText(
                    f"Impossibile visualizzare il PDF:\n{Path(item.path).name}\n\n"
                    f"Dettaglio: {detail}\n\n"
                    "Prova: pip install --upgrade PyMuPDF"
                )
            else:
                self.image_label.setText(f"Impossibile visualizzare:\n{Path(item.path).name}")
            self.image_label.setPixmap(QPixmap())
        else:
            self._render_zoom()

    def preload(self, item: FileItem | None) -> None:
        """Precarica in cache la pixmap del prossimo file immagine o
        PDF, per zero tempi morti quando l'utente passa al file
        successivo. Gli altri tipi (video, testo, segnaposto) non
        vengono precaricati in questa cache. Anche il precaricamento
        avviene in un thread separato, per lo stesso motivo della
        visualizzazione: un file lento da decodificare non deve mai
        bloccare l'interfaccia, nemmeno "in anticipo"."""
        if item is None or item.path in self._cache:
            return
        if item.category not in (FileCategory.IMAGE, FileCategory.PDF):
            return
        worker = _PreviewRenderWorker(item, RENDER_MAX, self)
        worker.done.connect(self._on_preload_done)
        self._preload_workers.append(worker)  # riferimento tenuto vivo finche' non finisce
        worker.start()

    def _on_preload_done(self, path: str, qimage) -> None:
        worker = self.sender()
        if worker in self._preload_workers:
            self._preload_workers.remove(worker)
        if qimage is None:
            return
        self._cache[path] = QPixmap.fromImage(qimage)
        if len(self._cache) > 8:
            self._cache.pop(next(iter(self._cache)))

    # -------------------------------------------------------------- video
    def is_showing_video(self) -> bool:
        return self.stack.currentIndex() == 3

    def toggle_play_pause(self) -> None:
        """Metodo pubblico usato anche dalla scorciatoia barra
        spaziatrice quando e' visualizzato un video."""
        self._toggle_play_pause()

    def _on_big_play_clicked(self) -> None:
        self._toggle_play_pause()

    SPEED_LEVELS = (1.0, 2.0, 4.0, 8.0)

    def _cycle_playback_speed(self) -> None:
        if self.player is None:
            return
        current = self.player.playbackRate()
        try:
            idx = self.SPEED_LEVELS.index(current)
        except ValueError:
            idx = 0
        next_speed = self.SPEED_LEVELS[(idx + 1) % len(self.SPEED_LEVELS)]
        self.player.setPlaybackRate(next_speed)
        self.speed_btn.setText(f"{int(next_speed)}x")

    def _toggle_play_pause(self) -> None:
        if self.player is None:
            return
        if self.player.playbackState() == QMediaPlayer.PlayingState:
            self.player.pause()
        else:
            self.player.play()
            # appena parte, passa dalla pagina "poster+play" a quella
            # col video vero: da qui in poi resta li' (anche in pausa)
            self.video_inner_stack.setCurrentIndex(1)

    def _on_video_state_changed(self, state) -> None:
        if self.player is not None and state == QMediaPlayer.PlayingState:
            self.play_btn.setText("⏸")
        else:
            self.play_btn.setText("▶")

    def _on_video_position_changed(self, position_ms: int) -> None:
        self.video_slider.blockSignals(True)
        self.video_slider.setValue(position_ms)
        self.video_slider.blockSignals(False)
        duration = self.player.duration() if self.player else 0
        self.video_time_label.setText(f"{_format_ms(position_ms)} / {_format_ms(duration)}")

    def _on_video_duration_changed(self, duration_ms: int) -> None:
        self.video_slider.setRange(0, max(0, duration_ms))

    # -------------------------------------------------------------- zoom
    def _render_zoom(self) -> None:
        if self._base_pixmap is None:
            return
        viewport = self.scroll.viewport().size()
        fit_scale = min(
            viewport.width() / max(self._base_pixmap.width(), 1),
            viewport.height() / max(self._base_pixmap.height(), 1),
            1.0,  # non ingrandire oltre la risoluzione nativa quando zoom=1
        )
        scale = fit_scale * self._zoom
        target = self._base_pixmap.size() * scale
        scaled = self._base_pixmap.scaled(
            max(target.width(), 1), max(target.height(), 1),
            Qt.KeepAspectRatio, Qt.SmoothTransformation,
        )
        self.image_label.setText("")
        self.image_label.setPixmap(scaled)
        self.image_label.resize(scaled.size())
        self._update_recenter_visibility()

    def _handle_wheel_zoom(self, event) -> bool:
        """Applica lo zoom in base alla rotellina, mantenendo fermo
        sotto al cursore il punto dell'immagine su cui si trova (zoom
        'ancorato al mouse', come nella maggior parte dei visualizzatori):
        si calcola su quale punto dell'immagine si trova il cursore
        PRIMA dello zoom, si applica lo zoom, poi si sposta la vista in
        modo che quello stesso punto resti sotto al cursore."""
        if self.stack.currentIndex() != 0 or self._base_pixmap is None:
            return False
        delta = event.angleDelta().y()
        if delta == 0:
            return False

        h_bar = self.scroll.horizontalScrollBar()
        v_bar = self.scroll.verticalScrollBar()
        viewport_pos = self.scroll.viewport().mapFromGlobal(event.globalPosition().toPoint())

        old_w = max(self.image_label.width(), 1)
        old_h = max(self.image_label.height(), 1)
        # punto dell'immagine sotto al cursore, come frazione (0-1) delle dimensioni attuali
        frac_x = (h_bar.value() + viewport_pos.x()) / old_w
        frac_y = (v_bar.value() + viewport_pos.y()) / old_h

        if delta > 0:
            self._zoom = min(self._zoom * self.ZOOM_STEP, self.ZOOM_MAX)
        else:
            self._zoom = max(self._zoom / self.ZOOM_STEP, self.ZOOM_MIN)

        self._render_zoom()

        new_w = max(self.image_label.width(), 1)
        new_h = max(self.image_label.height(), 1)
        h_bar.setValue(int(frac_x * new_w - viewport_pos.x()))
        v_bar.setValue(int(frac_y * new_h - viewport_pos.y()))

        event.accept()
        return True

    def wheelEvent(self, event) -> None:
        if not self._handle_wheel_zoom(event):
            super().wheelEvent(event)

    def _on_scroll_wheel(self, event) -> None:
        """La rotellina, quando passa sopra l'area di scorrimento
        dell'immagine, deve continuare a controllare lo zoom anche
        dopo che l'immagine e' diventata piu' grande della finestra
        (altrimenti l'area di scorrimento la 'ruba' per far scorrere
        invece di zoomare, ed e' impossibile tornare indietro)."""
        if not self._handle_wheel_zoom(event):
            QScrollArea.wheelEvent(self.scroll, event)

    def reset_zoom(self) -> None:
        self._zoom = 1.0
        self._render_zoom()

    # ------------------------------------------------------------ layout
    def _position_counter(self) -> None:
        margin = 12
        x = self.width() - self.counter_label.width() - margin
        y = self.height() - self.counter_label.height() - margin
        self.counter_label.move(max(0, x), max(0, y))

    def _position_type_badge(self) -> None:
        margin = 12
        x = self.width() - self.type_badge_label.width() - margin
        self.type_badge_label.move(max(0, x), margin)

    def _position_recenter_btn(self) -> None:
        margin = 12
        self.recenter_btn.move(margin, self.height() - self.recenter_btn.height() - margin)

    def _update_recenter_visibility(self) -> None:
        show = abs(self._zoom - 1.0) > 0.01 and self.stack.currentIndex() == 0
        if show:
            self._position_recenter_btn()
            self.recenter_btn.show()
            self.recenter_btn.raise_()
        else:
            self.recenter_btn.hide()

    def _update_type_badge(self, item: FileItem | None) -> None:
        """Mostra l'icona TUF dell'estensione in alto a destra per i
        tipi NON multimediali (documenti/archivi/codice/3D/audio/
        ebook/sottotitoli/testo/font/applicazioni), grande circa 1/16
        dell'area del riquadro. Le foto/video/GoPro-360-Drone non la
        mostrano, dato che gia' si vede (o si prova a mostrare) il
        contenuto vero."""
        if item is None or item.category in _NO_BADGE_CATEGORIES:
            self.type_badge_label.hide()
            return
        pix = _load_badge_pixmap(item.extension)
        if pix is None:
            self.type_badge_label.hide()
            return
        side = max(28, min(self.width(), self.height()) // 4)  # area ~= side^2 = (min/4)^2 ~ 1/16 del riquadro
        self.type_badge_label.setFixedSize(side, side)
        radius = max(4, side // 5)
        self.type_badge_label.setPixmap(_rounded_pixmap(pix, side, radius))
        self.type_badge_label.setStyleSheet("background-color: transparent;")
        self._position_type_badge()
        self.type_badge_label.show()
        self.type_badge_label.raise_()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._position_counter()
        self._position_type_badge()
        self._position_recenter_btn()
        if self._current_item is not None and self.stack.currentIndex() == 0:
            self._render_zoom()
